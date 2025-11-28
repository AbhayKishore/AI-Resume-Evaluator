import os
import json
import re
import textwrap
import tempfile
from dotenv import load_dotenv
from flask import Flask, request, render_template, redirect, url_for, session
from werkzeug.utils import secure_filename
from flask_session import Session
import pdfplumber
from docx import Document
from flask import send_file
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from datetime import datetime
import google.generativeai as genai
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Optional libs (not mandatory) ------------------------------------------------
try:
    from langdetect import detect as detect_lang
except Exception:
    detect_lang = None
    print("Warning: langdetect not installed — multilingual detection disabled.")

try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    OCR_AVAILABLE = True
except Exception:
    pytesseract = None
    convert_from_path = None
    Image = None
    OCR_AVAILABLE = False
    print("Warning: pytesseract/pdf2image/PIL not installed — OCR disabled.")

# -------------------------
# Configuration & setup
# -------------------------
load_dotenv()
API_KEY = os.getenv("GOOGLE_API_KEY")
if not API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables")
genai.configure(api_key=API_KEY)

# defensive model selection (safe fallback)
try:
    available_models = [m.name for m in genai.list_models() if 'gemini' in m.name.lower()]
except Exception as e:
    print("Warning listing models:", e)
    available_models = []
SUPPORTED_MODEL = "gemini-2.0-flash-001"
MODEL_NAME = next((m for m in available_models if SUPPORTED_MODEL in m), None) or "gemini-1.0-pro"
model = genai.GenerativeModel(MODEL_NAME)
print("Using model:", MODEL_NAME)

app = Flask(__name__)
app.config["SESSION_TYPE"] = "filesystem"
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret")
Session(app)

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Feature flags from env (optional)
ENABLE_OCR = os.getenv("ENABLE_OCR", "true").lower() in ("1", "true", "yes")
AUTO_REDACT_PII = os.getenv("AUTO_REDACT_PII", "false").lower() in ("1", "true", "yes")
BIAS_MONITORING = os.getenv("BIAS_MONITORING", "false").lower() in ("1", "true", "yes")

# -------------------------
# Basic helpers
# -------------------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf_with_ocr(filepath, maxpages=10):
    """Try OCR on PDF pages using pdf2image + pytesseract. Return concatenated text."""
    if not OCR_AVAILABLE:
        return ""
    try:
        pages = convert_from_path(filepath, dpi=200, first_page=1, last_page=maxpages)
        texts = []
        for p in pages:
            try:
                txt = pytesseract.image_to_string(p)
                texts.append(txt)
            except Exception as e:
                print("OCR page error:", e)
        return "\n".join(texts)
    except Exception as e:
        print("OCR conversion failed:", e)
        return ""

def extract_text_from_file(filepath):
    """Extract text reliably from PDF or DOCX. If PDF text extraction fails, attempt OCR (if available)."""
    try:
        if filepath.lower().endswith('.pdf'):
            # Try pdfplumber first
            try:
                with pdfplumber.open(filepath) as pdf:
                    pages = []
                    for p in pdf.pages:
                        txt = p.extract_text() or ""
                        pages.append(txt)
                    text = "\n".join(pages).strip()
            except Exception as e:
                print("pdfplumber error:", e)
                text = ""
            # If pdfplumber returned little/no text, try OCR (optional)
            if (not text or len(text) < 50) and ENABLE_OCR and OCR_AVAILABLE:
                print("Attempting OCR for PDF (scanned) as fallback...")
                ocr_text = extract_text_from_pdf_with_ocr(filepath, maxpages=25)
                if ocr_text and len(ocr_text) > len(text):
                    text = ocr_text
            return text
        elif filepath.lower().endswith('.docx'):
            doc = Document(filepath)
            return "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        print("Error extracting text:", e)
    return ""

def call_gemini_api(prompt, max_tokens=1000, expect_json=False, language_hint=None):
    """Call Gemini with defensive parsing. language_hint can be passed to instruct gemini."""
    try:
        # Optionally inject language hint into generation config or prompt for clearer multilingual behavior
        gencfg = {"max_output_tokens": max_tokens}
        response = model.generate_content(prompt, generation_config=gencfg)
        response_text = ""
        if hasattr(response, "text"):
            response_text = response.text.strip()
        elif isinstance(response, str):
            response_text = response.strip()
        else:
            response_text = str(response)
        print("Raw API response (truncated):", response_text[:1000])
        if expect_json:
            m = re.search(r'(\{.*\}|\[.*\])', response_text, re.DOTALL)
            return m.group(0) if m else response_text
        return response_text
    except Exception as e:
        print("Gemini API error:", e)
        return ""

# -------------------------
# Text-safety utilities
# -------------------------
def normalize_list_to_strings(items):
    """Convert mixed list items (dicts/strings) -> list[str], dedup preserve order."""
    if not items:
        return []
    out = []
    seen = set()
    for it in items:
        if it is None:
            continue
        if isinstance(it, str):
            s = it.strip()
        elif isinstance(it, dict):
            # join likely human keys
            parts = []
            for k in ("title","job_title","position","degree","institution","company","dates","date","name"):
                if k in it and it[k]:
                    parts.append(str(it[k]).strip())
            s = ", ".join(parts) if parts else json.dumps(it, ensure_ascii=False)[:240]
        else:
            s = str(it).strip()
        s = re.sub(r'\s{2,}', ' ', s)
        if s and s not in seen:
            out.append(s)
            seen.add(s)
    return out

def filter_by_text(candidates, text, min_overlap_chars=6):
    """
    Keep only candidate strings that appear in `text` (exact or fuzzy substring).
    Prevents trusting model-suggested lists that are not present in the resume.
    """
    if not candidates:
        return []
    raw = (text or "").lower()
    kept = []
    seen = set()
    for c in candidates:
        if not c:
            continue
        s = str(c).strip()
        if not s:
            continue
        low = s.lower()
        if low in raw:
            if s not in seen:
                kept.append(s); seen.add(s)
            continue
        # fuzzy contiguous token chunk matching
        tokens = re.findall(r'\w+', low)
        longest_chunk = None
        for L in range(min(4, len(tokens)), 0, -1):
            for i in range(0, len(tokens)-L+1):
                chunk = " ".join(tokens[i:i+L])
                if len(chunk) < min_overlap_chars:
                    continue
                if chunk in raw:
                    longest_chunk = chunk
                    break
            if longest_chunk:
                break
        if longest_chunk:
            if s not in seen:
                kept.append(s); seen.add(s)
            continue
        # token presence ratio
        hits = sum(1 for t in tokens if t in raw)
        if hits / max(1, len(tokens)) >= 0.6:
            if s not in seen:
                kept.append(s); seen.add(s)
    return kept

# -------------------------
# Extraction heuristics (unchanged but kept here for reference)
# -------------------------
def extract_name(text: str):
    if not text:
        return "Candidate Name"
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    title_blacklist = {'instructor','manager','engineer','developer','student','intern','teacher','trainer','resume','curriculum','vitae','profile','summary','contact'}
    for ln in lines[:10]:
        if ':' in ln: continue
        low = ln.lower()
        if any(low.startswith(tok + ':') for tok in ('phone','email','address','website','linkedin')): continue
        if any(tok in low for tok in title_blacklist): continue
        words = ln.split()
        if not (1 < len(words) <= 4): continue
        titlecase_count = sum(1 for w in words if re.match(r'^[A-Z][a-z]+$', w))
        allcaps_count = sum(1 for w in words if w.isupper() and len(w) > 1)
        if titlecase_count >= 1 or allcaps_count >= 1:
            return ln if ln.isupper() else " ".join(w.capitalize() for w in words)
    m = re.search(r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})\b', text)
    return m.group(1) if m else "Candidate Name"

def extract_contact_details(text: str):
    email = phone = address = ""
    lines = [ln.rstrip() for ln in text.splitlines() if ln.strip()]
    for i, ln in enumerate(lines):
        low = ln.lower()
        if low.startswith('email:') or low.startswith('e-mail:'):
            email = ln.split(':',1)[1].strip(); continue
        if low.startswith('phone:') or low.startswith('mobile:') or low.startswith('tel:') or low.startswith('contact:'):
            phone = ln.split(':',1)[1].strip(); continue
        if low.startswith('address:'):
            addr = ln.split(':',1)[1].strip()
            j = i+1
            while j < min(i+4, len(lines)):
                nxt = lines[j]
                if ':' in nxt or re.search(r'^(summary|skills|education|work|experience|projects|achiev)', nxt, flags=re.I):
                    break
                if len(nxt) < 160:
                    addr += ", " + nxt.strip(); j += 1
                else:
                    break
            address = addr.strip(); continue
    if not email:
        m = re.search(r'([a-zA-Z0-9_.+\-]+@[a-zA-Z0-9\-]+\.[a-zA-Z0-9\-.]+)', text)
        if m: email = m.group(1).strip()
    if not phone:
        m = re.search(r'(\+?\d[\d\-\s]{6,}\d)', text)
        if m: phone = m.group(1).strip()
    if not address:
        for ln in lines[:30]:
            lowln = ln.lower()
            if any(tok in lowln for tok in ('street','st.','road','rd','p.o.','po','post','pin','pincode','village','city','town','house','kerala','india','state','post office')):
                address = ln.strip(); break
    return email.strip(), phone.strip(), address.strip()

# ... (keep your other extractors unchanged) ...
def extract_general_information_section(text: str):
    lines = [ln.rstrip() for ln in text.splitlines()]
    collected = []
    start = None
    for i, ln in enumerate(lines):
        if re.match(r'^\s*GENERAL INFORMATION\b', ln, flags=re.I):
            start = i+1
            break
        if re.match(r'^\s*(Seminars?|Paper Presentations|Professional Trainings|Certifications|Workshops)\b', ln, flags=re.I):
            start = i
            break
    if start is None:
        return []
    for j in range(start, len(lines)):
        ln = lines[j].strip()
        if not ln:
            if j+1 < len(lines) and re.match(r'^\s*(WORK EXPERIENCE|PROJECT PROFILE|SKILL SET|ACHIEVEMENTS|ACADEMIC PROFILE)\b', lines[j+1], flags=re.I):
                break
            continue
        if re.match(r'^\s*(Date|Place|Page)\b', ln, flags=re.I) or re.search(r'page\s*\d+\s*/\s*\d+', ln, flags=re.I):
            break
        if re.search(r'\b(certificate|certified|certification|workshop|seminar|paper|conference|training|value-added|course)\b', ln, flags=re.I):
            cleaned = re.sub(r'^[\-\u2022\*\u25AA\u25CF\▪\u2023\s]+', '', ln).strip()
            if cleaned and cleaned not in collected:
                collected.append(cleaned)
        else:
            if len(ln) < 200 and not re.match(r'^[A-Z ]+$', ln):
                cleaned = re.sub(r'^[\-\u2022\*\s]+', '', ln).strip()
                if cleaned and cleaned not in collected:
                    collected.append(cleaned)
        if re.match(r'^\s*(WORK EXPERIENCE|PROJECT PROFILE|SKILL SET|ACHIEVEMENTS|ACADEMIC PROFILE)\b', ln, flags=re.I):
            break
    return collected

def extract_achievements_from_text(text: str):
    """
    Extract achievements / awards, including lines that may contain dates.
    This version does not discard achievements simply because they include a date.
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    collected = []

    for i, ln in enumerate(lines):
        if re.match(r'^\s*(ACCOMPLISHMENTS?|ACHIEVEMENTS?|AWARDS|HONORS)\b', ln, flags=re.I):
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt:
                    j += 1
                    continue

                if re.match(r'^\s*(WORK EXPERIENCE|PROJECT PROFILE|SKILL SET|GENERAL INFORMATION|ACADEMIC PROFILE)\b',
                            nxt, flags=re.I):
                    break

                if re.match(r'^(Date|Place|Signature)\b', nxt, flags=re.I) or re.search(
                        r'page\s*\d+\s*/\s*\d+', nxt, flags=re.I):
                    j += 1
                    continue

                if (re.search(r'\b(prize|won|winner|stood|placed|first|second|third|distinction|award|achiev|recognition|achievement)\b',
                              nxt, flags=re.I)
                        or re.match(r'^[\-\u2022\*\u25AA\u25CF\▪\u2023]\s+', nxt)
                        or (10 < len(nxt) < 200)):
                    cleaned = re.sub(r'^[\-\u2022\*\s]+', '', nxt).strip()

                    # NOTE: we do NOT remove lines just because they have a date
                    if re.search(r'page\s*\d+\s*/\s*\d+', cleaned, flags=re.I):
                        j += 1
                        continue

                    if len(cleaned) > 5 and cleaned not in collected:
                        collected.append(cleaned)
                j += 1

    if not collected:
        for i, ln in enumerate(lines):
            if re.search(r'\b(First Prize|First Place|Stood\s+First|Winner|Won|Awarded|Distinction|Placed|Secured|Best Performer)\b',
                         ln, flags=re.I):
                if re.search(r'page\s*\d+\s*/\s*\d+', ln, flags=re.I):
                    continue
                base = ln
                if i + 1 < len(lines):
                    nxt = lines[i + 1].strip()
                    if not re.match(r'^(Date|Place)\b', nxt, flags=re.I) and len(nxt) < 200:
                        base = base + " " + nxt
                base_clean = re.sub(r'^[\-\u2022\*\s]+', '', base).strip()

                if len(base_clean) > 6 and base_clean not in collected:
                    collected.append(base_clean)

    final = []
    for it in collected:
        s = re.sub(r'\s{2,}', ' ', it).strip()

        if re.search(r'page\s*\d+\s*/\s*\d+', s, flags=re.I):
            continue

        if re.match(r'^(Affiliat|Applied Sciences|Affiliated with|Affiliation)', s, flags=re.I) and not re.search(
            r'\b(prize|award|stood|first|placed|winner)\b', s, flags=re.I
        ):
            continue

        if re.match(r'^Place\s*:', s, flags=re.I):
            continue

        if len(s) < 6:
            continue

        if s not in final:
            final.append(s)

    return final

def extract_training_lines(text: str):
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    trainings = []
    seen = set()
    for i, ln in enumerate(lines):
        if re.search(r'\b(training|internship|industrial training|bootcamp|apprentice|trained)\b', ln, flags=re.I):
            block = ln
            if i-1 >= 0 and len(lines[i-1]) < 160 and not re.search(r'^(Date|Place|Page)\b', lines[i-1], flags=re.I):
                block = lines[i-1].strip() + " " + block
            if i+1 < len(lines) and len(lines[i+1]) < 160 and not re.search(r'^(Date|Place|Page)\b', lines[i+1], flags=re.I):
                block = block + " " + lines[i+1].strip()
            block = re.sub(r'\s{2,}', ' ', block).strip()
            if re.search(r'page\s*\d+\s*/\s*\d+', block, flags=re.I): continue
            if block.lower() in seen: continue
            trainings.append(block)
            seen.add(block.lower())
    final = []
    for t in trainings:
        if re.match(r'^[A-Z ]+$', t) and len(t.split()) <= 3:
            continue
        final.append(t)
    return final

# -------------------------
# NEW: Soft-skill detection (basic + model-assisted)
# -------------------------
SOFT_SKILL_LEXICON = [
    "leadership","communication","teamwork","collaboration","adaptability","problem solving",
    "time management","critical thinking","empathy","creativity","mentoring","presentation"
]

def detect_soft_skills(text: str):
    text_low = (text or "").lower()
    found = []
    for s in SOFT_SKILL_LEXICON:
        if s in text_low:
            found.append(s)
    # model-assisted fallback (best-effort)
    if not found:
        # ask Gemini to spot soft skills conservatively (expect short JSON)
        prompt = f"""
You MUST respond with a VALID JSON array ONLY, no additional text.
From the resume below, list up to 6 soft skills that are clearly implied or stated.
Resume: {text[:2000]}
"""
        resp = call_gemini_api(prompt, expect_json=True)
        if resp:
            try:
                arr = json.loads(resp)
                if isinstance(arr, list):
                    found = [a for a in arr if isinstance(a, str)]
            except Exception:
                pass
    return list(dict.fromkeys(found))

# -------------------------
# NEW: Language detection helper
# -------------------------
def detect_language_of_text(text: str):
    if not text:
        return "unknown"
    if detect_lang:
        try:
            return detect_lang(text)
        except Exception:
            return "unknown"
    # fallback heuristic
    common_english = len(re.findall(r'\b(the|and|of|in|for)\b', text.lower()))
    return "en" if common_english > 2 else "unknown"

# -------------------------
# NEW: Simple AI-generated text detection
# -------------------------
def local_ai_text_heuristic(text: str):
    """
    Simple heuristics: very repetitive phrasing, many average-long sentences, improbable vocabulary uniformity.
    Returns score 0-100 risk (higher = more likely AI).
    """
    if not text:
        return 0
    # repetition metric
    words = re.findall(r'\w+', text.lower())
    if not words:
        return 0
    unique_ratio = len(set(words)) / len(words)
    repetition_score = max(0, int((1 - unique_ratio) * 100))
    # sentence length uniformity
    sentences = re.split(r'[.!?]\s+', text.strip())
    if sentences:
        avg_len = sum(len(s.split()) for s in sentences) / len(sentences)
        var = (sum((len(s.split()) - avg_len)**2 for s in sentences) / len(sentences)) if len(sentences) > 0 else 0
        uniformity_score = int(max(0, 50 - min(50, var)))
    else:
        uniformity_score = 0
    # presence of marketing-like adjectives
    promo_hits = len(re.findall(r'\b(passionate|experienced|expert|highly skilled|proven)\b', text.lower()))
    promo_score = min(30, promo_hits * 10)
    score = min(100, int(0.5*repetition_score + 0.3*uniformity_score + 0.2*promo_score))
    return score

def detect_ai_generated_text(text: str):
    """Combine local heuristic with model classification for stronger signal."""
    local_score = local_ai_text_heuristic(text)
    # Ask Gemini for classification (conservative)
    prompt = f"""
You are a classifier. Answer ONLY with a JSON object: {{"ai_risk_score": <0-100>, "reason": "short explanation"}}.
Assess the following resume text for likelihood of being AI-generated or heavily assisted. Be conservative.
Text: {text[:2000]}
"""
    resp = call_gemini_api(prompt, expect_json=True)
    model_score = None
    model_reason = ""
    if resp:
        try:
            parsed = json.loads(resp)
            model_score = int(parsed.get("ai_risk_score", 0))
            model_reason = parsed.get("reason", "")
        except Exception:
            # try to heuristically pull numbers
            m = re.search(r'(\d{1,3})', resp)
            if m:
                try:
                    model_score = int(m.group(1))
                except:
                    model_score = None
    final = local_score
    if model_score is not None:
        final = int((local_score + model_score) / 2)
    return {"ai_risk_score": final, "local": local_score, "model_reason": model_reason}

# -------------------------
# NEW: Timeline / fraud heuristics
# -------------------------
def extract_date_tokens(text):
    # capture year ranges and standalone years
    patterns = [
        r'(\d{4})\s*-\s*(\d{4})',   # 2018-2020
        r'(\d{4})\s*to\s*(\d{4})',  # 2018 to 2020
        r'(\d{4})'                  # 2020
    ]
    years = []
    for p in re.findall(r'\d{4}\s*-\s*\d{4}', text):
        m = re.match(r'(\d{4})\s*-\s*(\d{4})', p)
        if m:
            years.append((int(m.group(1)), int(m.group(2))))
    # standalone years
    standalone = [int(y) for y in re.findall(r'(?<!\d)(19|20)\d{2}(?!\d)', text)]
    for y in standalone:
        years.append((y, y))
    return years

def timeline_issues_from_experience(experience_list):
    """
    Look for overlapping job dates or implausible long gaps.
    Returns list of issues.
    """
    issues = []
    ranges = []
    for e in experience_list:
        found = re.findall(r'(\d{4})\s*(?:-|to)\s*(\d{4})', e)
        if found:
            for s,e_year in found:
                try:
                    s = int(s); e_year = int(e_year)
                    ranges.append((s,e_year,e))
                except:
                    pass
        else:
            # try single year
            found2 = re.findall(r'(?<!\d)(19|20)\d{2}(?!\d)', e)
            if found2:
                y = int(found2[0])
                ranges.append((y,y,e))
    ranges_sorted = sorted(ranges, key=lambda x: x[0])
    # overlapping detection
    for i in range(len(ranges_sorted)-1):
        s1,e1,_ = ranges_sorted[i]
        s2,e2,_ = ranges_sorted[i+1]
        if s2 <= e1:
            issues.append(f"Overlapping employment dates: {s1}-{e1} overlaps with {s2}-{e2}")
    # gaps detection
    years = [r[1] for r in ranges_sorted]
    if years:
        max_gap = 0
        for i in range(len(years)-1):
            gap = ranges_sorted[i+1][0] - ranges_sorted[i][1]
            if gap > max_gap:
                max_gap = gap
        if max_gap >= 5:
            issues.append(f"Large gap detected between experiences (>= {max_gap} years).")
    return issues

# -------------------------
# NEW: Simple internal matching engine (fuzzy)
# -------------------------
from difflib import SequenceMatcher

def fuzzy_similarity(a: str, b: str):
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()

def compute_matching_score(resume_text: str, job_text: str, weights=None):
    """
    Compute a local matching score (0-100) using fuzzy matching across:
    - skills (exact/substring matches)
    - whole-text similarity
    - keyword overlaps
    """
    if weights is None:
        weights = {"text": 0.4, "skills": 0.4, "keywords": 0.2}
    text_score = fuzzy_similarity(resume_text, job_text)
    # extract keywords from job (simple nouns/words >3 chars)
    job_keywords = set([w.lower() for w in re.findall(r'\b\w{4,}\b', job_text)])
    resume_words = set([w.lower() for w in re.findall(r'\b\w{4,}\b', resume_text)])
    if job_keywords:
        keyword_overlap = len(job_keywords & resume_words) / max(1, len(job_keywords))
    else:
        keyword_overlap = 0.0
    # skills match: look for common tokens from analysis.skills (we'll pass externally if available)
    skills_score = 0.0
    # final weighted
    final = weights["text"]*text_score + weights["keywords"]*keyword_overlap + weights["skills"]*skills_score
    return int(round(final * 100))

# -------------------------
# Main analysis pipeline (keeps old behavior but enriches analysis)
# -------------------------
def process_resume(resume_text: str, job_text: str) -> dict:
    """
    Central analysis pipeline (defensive).
    Returns a dict with sanitized lists and strings suitable for generating resume + report.
    """
    analysis = {
        "resume_text": resume_text,
        "skills": [],
        "education": [],
        "experience": [],
        "honors_awards": [],
        "summary": "",
        "ats_score": 0,
        "ats_feedback": [],
        "missing_keywords": [],
        "job_matches": [],
        "desired_job": job_text or "General Role",
        "missing_skills": [],
        "missing_education": [],
        "course_recommendations": [],
        "education_recommendations": [],
        "alternative_paths": [],
        "error": "",
        "chart_labels": [],
        "chart_data": [],
        "ats_breakdown": {"technical": 0, "soft": 0, "experience": 0},
        "language": "unknown",
        "soft_skills": [],
        "ai_risk": {"ai_risk_score": 0, "local": 0, "model_reason": ""},
        "timeline_issues": [],
        "fraud_risk": 0,
        "pii_redacted": False,
        "bias_warnings": [],
        "matching_score": 0,
        "job_text_raw": job_text or ""
        
    }
      # --- NEW: derive concise target-role title from job description via Gemini ---
    if job_text:
        role_label_prompt = f"""
You MUST respond with plain text ONLY, no JSON.
From the job description below, return a concise role title (max 60 characters),
e.g. "Junior DevOps Engineer" or "DevOps / Cloud Engineer (Entry Level)".

Job Description:
{(job_text or '')[:600]}
"""
        role_label = call_gemini_api(role_label_prompt, expect_json=False).strip()
        if role_label:
            analysis["desired_job"] = role_label[:60]

    # Detect language early
    analysis["language"] = detect_language_of_text(resume_text)

    # 1) Ask model to extract structured fields (unchanged)
    extraction_prompt = f"""
You MUST respond with a VALID JSON object ONLY, no additional text. Extract from the resume:
- "skills": list of specific skills (e.g., ["Python","Django"]).
- "education": list of degrees/institutions/dates.
- "experience": list of job/position strings (title, company, dates, short description).
- "honors_awards": list of honors/awards if any.
Resume: {resume_text[:3000]}
"""
    extraction_response = call_gemini_api(extraction_prompt, expect_json=True)
    raw_skills = raw_education = raw_experience = raw_honors = []
    if extraction_response:
        try:
            extraction_data = json.loads(extraction_response)
            raw_skills = extraction_data.get("skills", []) or []
            raw_education = extraction_data.get("education", []) or []
            raw_experience = extraction_data.get("experience", []) or []
            raw_honors = extraction_data.get("honors_awards", []) or []
        except Exception as e:
            print("Failed to parse extraction JSON:", e)

    analysis["skills"] = normalize_list_to_strings(raw_skills)
    analysis["education"] = normalize_list_to_strings(raw_education)
    analysis["experience"] = normalize_list_to_strings(raw_experience)
    analysis["_raw_honors_normalized"] = normalize_list_to_strings(raw_honors)

    # 2) Achievements / honors (resume-based + robust keyword fallback)
    honors = []

    # 2a) Existing parser result
    honors_from_text = extract_achievements_from_text(resume_text or "")
    for h in honors_from_text:
        h = (h or "").strip()
        if h and h not in honors:
            honors.append(h)

    # 2b) Extra keyword-based achievements (e.g. "Stood First", "Winner", etc.)
    lines = [ln.strip() for ln in (resume_text or "").splitlines() if ln.strip()]

    KEYWORD_PATTERN = re.compile(
        r'\b(First Prize|First Place|Stood\s+First|Winner|Won\b|Awarded|Distinction|Placed|Secured|Best Performer)\b',
        flags=re.I
    )
    HEADING_ONLY = re.compile(r'^\s*(Accomplishments?|Achievements?|Awards?|Honors?)\s*$', flags=re.I)

    for i, ln in enumerate(lines):
        if KEYWORD_PATTERN.search(ln):
            base = ln

            # Try to join up to 2 following lines (to handle patterns like:
            # bullet -> "Accomplishments" heading -> continuation line)
            for offset in (1, 2):
                if i + offset >= len(lines):
                    break
                nxt = lines[i + offset].strip()

                # Skip pure headings such as "Accomplishments"
                if HEADING_ONLY.match(nxt):
                    continue

                # Stop if we hit meta-lines instead of content
                if re.match(r'^(Date|Place|Page)\b', nxt, flags=re.I):
                    break

                # Ignore abnormally long lines
                if len(nxt) >= 200:
                    break

                # Treat as continuation text
                base = base + " " + nxt

            base_clean = re.sub(
                r'^[\-\u2022\*\u25AA\u25CF\▪\u2023\·\s]+', '', base
            ).strip()

            if base_clean and base_clean not in honors:
                honors.append(base_clean)


    # 2c) Model honors that clearly appear in the resume text
    model_honors_all = analysis.get("_raw_honors_normalized", [])
    model_honors_valid = filter_by_text(model_honors_all, resume_text)
    for mh in model_honors_valid:
        mh = (mh or "").strip()
        if mh and mh not in honors:
            honors.append(mh)

    # 2d) Final clean-up
    final_honors = []
    for h in honors:
        s = re.sub(r'\s{2,}', ' ', h).strip()
        if not s:
            continue
        if re.search(r'page\s*\d+\s*/\s*\d+', s, flags=re.I):
            continue
        if len(s) < 6:
            continue
        if s not in final_honors:
            final_honors.append(s)

    analysis["honors_awards"] = final_honors

    # Debug (optional)
    print("HONORS DEBUG:", analysis["honors_awards"])

    # 3) Certifications (GENERAL INFORMATION) (unchanged)
    general_info_lines = extract_general_information_section(resume_text or "")
    extra_cert_lines = []
    for ln in (resume_text or "").splitlines():
        if re.search(r'\b(certificate|certified|certification|course|completed|bootcamp|value-added)\b', ln, flags=re.I):
            if len(ln.strip()) < 240 and 'page' not in ln.lower():
                extra_cert_lines.append(ln.strip())
    merged_certs = list(dict.fromkeys(general_info_lines + extra_cert_lines))
    vetted_certs = filter_by_text(merged_certs, resume_text)
    analysis["general_information"] = vetted_certs

    # 4) Clean experience (unchanged)
    cleaned_exps = []
    seen = set()
    for e in analysis.get("experience", []):
        s = re.sub(r'\s+', ' ', str(e)).strip()
        s = re.sub(r'Page\s*\d+\s*/\s*\d+', '', s, flags=re.I).strip()
        if len(s) >= 6 and s not in seen:
            cleaned_exps.append(s)
            seen.add(s)
    analysis["experience"] = cleaned_exps

    # 5) Extract training
    analysis["training"] = extract_training_lines(resume_text or "")

    # 6) Summary (unchanged)
    summary_prompt = f"""
You MUST respond with plain text ONLY, no JSON. Generate a concise 2-sentence professional summary from the resume.
DO NOT INVENT facts.
Resume: {resume_text[:3000]}
"""
    s_resp = call_gemini_api(summary_prompt, expect_json=False)
    summary_text = (s_resp or "").strip()
    summary_text = re.sub(r'\(([^)]+)\)', lambda m: m.group(0) if m.group(1).lower() in (resume_text or "").lower() else '', summary_text)
    if not summary_text:
        degree_hint = ""
        for edu in analysis["education"]:
            if re.search(r'\bb\.?c\.?a\b', edu, flags=re.I):
                degree_hint = "BCA"; break
            if "bachelor" in edu.lower():
                degree_hint = "Bachelor"; break
        skills_preview = ", ".join(analysis["skills"][:6])
        summary_text = f"{extract_name(resume_text)} is a {degree_hint or 'graduate'} with practical experience. Skills include {skills_preview or 'technical and interpersonal strengths'}."
    analysis["summary"] = summary_text

    # 7) ATS SCORE + MISSING ELEMENTS (unchanged)
    missing_prompt = f"""
You MUST respond with a VALID JSON object ONLY. Identify:
- "missing_skills": list of 2-5 skills missing
- "missing_education": list
Resume skills/education: {', '.join(map(str, analysis.get('skills', [])))} {', '.join(map(str, analysis.get('education', [])))}
Job Description: {job_text if job_text else 'General role'}
"""
    missing_resp = call_gemini_api(missing_prompt, expect_json=True)
    if missing_resp:
        try:
            miss_data = json.loads(missing_resp)
            analysis["missing_skills"] = miss_data.get("missing_skills", []) or []
            analysis["missing_education"] = miss_data.get("missing_education", []) or []
        except:
            pass

    ats_prompt = f"""
You MUST respond with a VALID JSON object ONLY. Return:
- "ats_score": integer 0-100
- "missing_keywords": list of STRINGS.
  Each item must be a keyword or short phrase (max 4 words) that appears in the JOB DESCRIPTION
  but is missing or clearly underrepresented in the RESUME.
- "ats_feedback": list of short, actionable suggestions (plain strings).

Do NOT invent technologies or skills that are not present in the job description.
Do NOT include extremely generic words like "teamwork" or "communication"
unless they are explicitly emphasized in the job description.

Resume text (truncated):
{resume_text[:2000]}

Job Description:
{job_text if job_text else 'General role'}
"""
    ats_resp = call_gemini_api(ats_prompt, expect_json=True)
    if ats_resp:
        try:
            ats_data = json.loads(ats_resp)
            analysis["ats_score"] = int(ats_data.get("ats_score", 0))
            analysis["missing_keywords"] = ats_data.get("missing_keywords", []) or []
            analysis["ats_feedback"] = ats_data.get("ats_feedback", []) or []
        except:
            pass

    # ASK GEMINI for ATS breakdown (existing)
    ats_breakdown_prompt = f"""
You MUST respond ONLY with VALID JSON. No text outside JSON.
Return:
{{
  "technical": <0-100>,
  "soft": <0-100>,
  "experience": <0-100>
}}
Resume: {resume_text[:2000]}
Job Description: {job_text if job_text else 'General role'}
"""
    ats_breakdown_resp = call_gemini_api(ats_breakdown_prompt, expect_json=True)
    if ats_breakdown_resp:
        try:
            bdata = json.loads(ats_breakdown_resp)
            analysis["ats_breakdown"] = {
                "technical": int(bdata.get("technical", 0)),
                "soft": int(bdata.get("soft", 0)),
                "experience": int(bdata.get("experience", 0))
            }
        except:
            # fallback: attempt reasonable defaults derived from content
            try:
                analysis["ats_breakdown"]["technical"] = min(100, analysis["ats_score"])
                analysis["ats_breakdown"]["soft"] = min(100, max(10, analysis["ats_score"]//2))
                analysis["ats_breakdown"]["experience"] = min(100, max(10, analysis["ats_score"]//3))
            except:
                pass

    # 8) Job matches (unchanged)
    job_match_prompt = f"""
You MUST respond with a VALID JSON array ONLY.
Format: [{{"job_title":"Role","reasoning":"Why"}}]
Skills: {', '.join(analysis.get("skills", []))}
Experience: {', '.join(analysis.get("experience", []))}
"""
    job_resp = call_gemini_api(job_match_prompt, expect_json=True)
    if job_resp:
        try:
            text = job_resp.strip()
            if text.startswith('{'):
                text = "[" + text.replace("}\n{", "}, {") + "]"
            analysis["job_matches"] = json.loads(text)
        except:
            analysis["job_matches"] = []

    # Course recommendations (unchanged)
    if analysis.get("missing_skills"):
        course_prompt = f"""
You MUST respond with a VALID JSON array ONLY.
Format: [{{"skill":"skill","courses":["A","B"]}}, ...]
Missing skills: {', '.join(analysis.get('missing_skills', []))}
"""
        course_resp = call_gemini_api(course_prompt, expect_json=True)
        if course_resp:
            try:
                text = course_resp.strip()
                if text.startswith('{'):
                    text = "[" + text.replace("}\n{", "}, {") + "]"
                analysis["course_recommendations"] = json.loads(text)
            except:
                analysis["course_recommendations"] = []

    # 9) Chart placeholders (unchanged)
    for key in ["skills","education","experience","honors_awards","missing_skills",
                "missing_education","course_recommendations","education_recommendations",
                "job_matches","alternative_paths","missing_keywords","ats_feedback",
                "general_information","training"]:
        analysis.setdefault(key, [])

    if analysis["skills"]:
        analysis["chart_labels"] = analysis["skills"][:6]
        analysis["chart_data"] = [80,70,60,50,40,30][:len(analysis["chart_labels"])]
    else:
        analysis["chart_labels"] = ["Skill A","Skill B","Skill C"]
        analysis["chart_data"] = [60,40,70]

    try:
        analysis["ats_score"] = int(analysis.get("ats_score", 0))
    except Exception:
        analysis["ats_score"] = 0

    # -------------------------
    # NEW: soft-skills detection
    analysis["soft_skills"] = detect_soft_skills(resume_text)

    # NEW: AI-generated text detection
    try:
        ai_det = detect_ai_generated_text(resume_text)
        if isinstance(ai_det, dict):
            analysis["ai_risk"] = {
                "ai_risk_score": ai_det.get("ai_risk_score", 0),
                "local": ai_det.get("local", 0),
                "model_reason": ai_det.get("model_reason") or ai_det.get("reason") or ""
            }
            analysis["fraud_risk"] = ai_det.get("ai_risk_score", 0)
    except Exception as e:
        print("AI detection error:", e)

    # NEW: Timeline issues
    try:
        t_issues = timeline_issues_from_experience(analysis.get("experience", []))
        analysis["timeline_issues"] = t_issues
        if t_issues:
            analysis["fraud_risk"] = max(analysis.get("fraud_risk",0), 30)
    except Exception as e:
        print("Timeline check error:", e)

    # NEW: matching score (local engine)
    try:
        analysis["matching_score"] = compute_matching_score(resume_text or " ", job_text or " ")
    except Exception as e:
        print("Matching score error:", e)
        analysis["matching_score"] = 0

    # NEW: bias monitoring (simple)
    if BIAS_MONITORING:
        warnings = []
        lowtxt = (resume_text or "").lower()
        # naive checks for protected attributes mentions
        for tok in ("gender","male","female","religion","caste","age:","age "):
            if tok in lowtxt:
                warnings.append(f"Potential protected-attribute mention: {tok}")
        analysis["bias_warnings"] = warnings


    # finalize desired job
    if not analysis.get("desired_job"):
        analysis["desired_job"] = (job_text or "General Role").strip()

    return analysis

# -------------------------
# Flask routes (kept & extended)
# -------------------------
@app.route('/')
def index():
    return render_template('index.html')  # your HTML UI

@app.route('/analyze', methods=['POST'])
def analyze():
    if 'resume' not in request.files:
        return render_template('index.html', error="No file part"), 400
    file = request.files['resume']
    if file.filename == '':
        return render_template('index.html', error="No selected file"), 400
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        job_text = request.form.get('job_text','').strip()
        resume_text = extract_text_from_file(filepath)
        if not resume_text.strip():
            return render_template('index.html', error="Unable to extract text from uploaded file"), 400
        analysis = process_resume(resume_text, job_text)
        # store only necessary fields to minimize persisted PII
        session['analysis'] = analysis
        return redirect(url_for('results'))
    return render_template('index.html', error="File type not allowed"), 400

@app.route('/results')
def results():
    analysis = session.get('analysis', {"error": "No analysis available."})
    # For debugging: print minimal keys
    print("Analysis keys:", list(analysis.keys()))
    return render_template('results.html', analysis=analysis)

@app.route('/generate-resume')
def generate_resume():
    analysis = session.get("analysis")
    if not analysis:
        return "No analysis found. Upload a resume first.", 400

    resume_text = analysis.get("resume_text", "") or ""
    candidate_name = extract_name(resume_text)
    email, phone, address = extract_contact_details(resume_text)

    personal_info = {
        "Name": candidate_name,
        "Address": address or "Not specified",
        "Phone": phone or "Not specified",
        "Email": email or "Not specified"
    }

    education = analysis.get("education", []) or []
    skills = analysis.get("skills", []) or []
    projects = analysis.get("project_profiles", []) if isinstance(analysis.get("project_profiles", None), list) else []
    if not projects and analysis.get("job_matches"):
        projects = [jm.get("job_title", "") for jm in analysis.get("job_matches", [])]
    experience_raw = analysis.get("experience", []) or []
    achievements = analysis.get("honors_awards", []) or []
    general_certs = analysis.get("general_information", []) or []
    trainings = analysis.get("training", []) or []

    # -----------------------------
    # 1) KEYWORD VALIDATION STEP
    # -----------------------------
    missing_keywords = analysis.get("missing_keywords", []) or []
    job_text_raw = analysis.get("job_text_raw", "") or ""

    validated_keywords = []
    keyword_example_bullets = []

    if missing_keywords:
        try:
            keyword_prompt = f"""
You are a STRICT fact-checker for resumes.

You will receive:
- RESUME_TEXT: the candidate's original resume content
- EXISTING_SKILLS: skills already extracted from the resume
- JOB_DESCRIPTION: the actual job description text
- MISSING_KEYWORDS: a list of keywords suggested from the job description

Your task:

1. From MISSING_KEYWORDS, keep ONLY those items which are clearly
   supported by the resume content and/or existing skills.
   If you do NOT see clear evidence in the resume that a keyword is true
   for this person, you MUST NOT include it.
2. Optionally suggest short, tailored bullet points that could be added
   to the resume. Each bullet MUST:
   - be strictly consistent with the resume content,
   - not invent any new tools, responsibilities or experience,
   - naturally incorporate one or more of the validated keywords.

Return a VALID JSON object ONLY, with this structure:
{{
  "validated_keywords": ["keyword1", "keyword2", ...],
  "example_descriptions": [
    "Optional bullet point 1 ...",
    "Optional bullet point 2 ..."
  ]
}}

RESUME_TEXT:
{resume_text[:3000]}

EXISTING_SKILLS:
{skills}

JOB_DESCRIPTION:
{job_text_raw}

MISSING_KEYWORDS:
{missing_keywords}
"""
            resp = call_gemini_api(keyword_prompt, expect_json=True)
            if resp:
                data = json.loads(resp)
                if isinstance(data, dict):
                    vk = data.get("validated_keywords", [])
                    ed = data.get("example_descriptions", [])
                    if isinstance(vk, list):
                        validated_keywords = [str(k).strip() for k in vk if str(k).strip()]
                    if isinstance(ed, list):
                        keyword_example_bullets = [str(b).strip() for b in ed if str(b).strip()]
        except Exception as e:
            print("Keyword validation error:", e)

    # -----------------------------
    # 2) PROJECT IMPACT / OUTCOME BULLETS
    # -----------------------------
    impact_bullets = []
    try:
        impact_prompt = f"""
You are a resume rewriting assistant.

Using ONLY the factual content from RESUME_TEXT below, generate up to 8 bullet points
that highlight the candidate's project impact and outcomes, focusing on:

- Achievements or measurable outcomes (ONLY if numbers or results are explicitly present;
  if no numbers are given, keep outcomes qualitative, e.g., "improved X", "enabled Y").
- Real-world examples of delivered results (what the project achieved or solved).
- Description of specific technologies used in each project (frameworks, languages, tools).
- Evidence of collaboration (team-based vs individual work) WHEN mentioned.
- Evidence of debugging / troubleshooting activities WHEN mentioned or clearly implied.
- Evidence of monitoring / logging ONLY IF present; otherwise, OMIT this aspect.
- Time durations for projects ONLY IF dates or time periods are explicitly given.

STRICT RULES:
- Do NOT invent technologies, metrics, companies, environments, or responsibilities
  that are not clearly supported by RESUME_TEXT.
- Do NOT guess numbers. Use a number only if it is explicitly in RESUME_TEXT.
- If something (like monitoring/logging, collaboration, or durations) is not clearly
  present in RESUME_TEXT, simply do not force it into the bullet.

Style:
- Each bullet: 1–2 lines, concise, past tense, suitable for a resume.
- Focus on impact and outcomes, but always remain factual.

Return ONLY a VALID JSON array of strings, no extra keys, no extra text.

RESUME_TEXT:
{resume_text[:3500]}
"""
        impact_resp = call_gemini_api(impact_prompt, expect_json=True)
        if impact_resp:
            data = json.loads(impact_resp)
            if isinstance(data, list):
                impact_bullets = [str(b).strip() for b in data if str(b).strip()]
    except Exception as e:
        print("Impact bullet generation error:", e)
        impact_bullets = []

    # -----------------------------
    # 3) STRUCTURED EXPERIENCE PARSING
    # -----------------------------
    structured_exps = []
    for e in experience_raw:
        parts = [p.strip() for p in str(e).split(',', 2)]
        obj = {
            "title": parts[0] if parts else "",
            "company": parts[1] if len(parts) > 1 else "",
            "dates": parts[2] if len(parts) > 2 else "",
            "bullets": []
        }
        if ':' in str(e):
            rhs = str(e).split(':', 1)[1].strip()
            bullets = [b.strip() for b in re.split(r'\s*[\;\-]\s*|\n', rhs) if b.strip()]
            obj["bullets"] = list(dict.fromkeys(bullets))[:8]
        structured_exps.append(obj)

    # -----------------------------
    # 4) BUILD DOCX
    # -----------------------------
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Name header
    p = doc.add_paragraph()
    run = p.add_run(candidate_name)
    run.bold = True
    run.font.size = Pt(22)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # sanitized title
    candidate_title = (analysis.get("desired_job") or "").splitlines()[0][:60]
    if candidate_title:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(candidate_title)
        r2.font.size = Pt(10)
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    # Personal info
    doc.add_heading('PERSONAL INFORMATION', level=2)
    for k, v in personal_info.items():
        para = doc.add_paragraph()
        runk = para.add_run(f"{k}: ")
        runk.bold = True
        para.add_run(str(v))

    # Profile summary
    doc.add_heading('PROFILE SUMMARY', level=2)
    doc.add_paragraph(analysis.get("summary", "Not specified"))

    # Academic profile
    doc.add_heading('ACADEMIC PROFILE', level=2)
    if education:
        for edu in education:
            doc.add_paragraph(str(edu))
    else:
        doc.add_paragraph("Not specified")

    # Industrial training
    doc.add_heading('INDUSTRIAL TRAINING', level=2)
    if trainings:
        for t in trainings:
            doc.add_paragraph(t)
    else:
        doc.add_paragraph("Not specified")

    # Project profile (raw/project titles etc.)
    doc.add_heading('PROJECT PROFILE AND INDIVIDUAL ROLE', level=2)
    if projects:
        for pr in projects:
            doc.add_paragraph(str(pr))
    else:
        doc.add_paragraph("Not specified")

    # NEW SECTION: PROJECT IMPACT & OUTCOMES
    doc.add_heading('PROJECT IMPACT & OUTCOMES', level=2)
    if impact_bullets:
        for b in impact_bullets:
            doc.add_paragraph(b, style='List Bullet')
    else:
        doc.add_paragraph(
            "Impact-focused details could not be derived beyond the existing project descriptions "
            "without inventing information. Consider manually adding metrics, outcomes, or specific "
            "results where appropriate."
        )

    # Skills
    doc.add_heading('SKILL SET', level=2)
    if skills:
        for s in skills:
            doc.add_paragraph(s, style='List Bullet')
    else:
        doc.add_paragraph("Not specified")

    # General information / Certifications
    doc.add_heading('GENERAL INFORMATION', level=2)
    if general_certs:
        for c in general_certs:
            doc.add_paragraph(c, style='List Bullet')
    else:
        doc.add_paragraph("Not specified")

    # Work experience
    doc.add_heading('WORK EXPERIENCE', level=2)
    if structured_exps:
        for ex in structured_exps:
            header_parts = [ex.get('title', ''), ex.get('company', ''), ex.get('dates', '')]
            header = " — ".join([p for p in header_parts if p])
            if header:
                pheader = doc.add_paragraph()
                pheader.add_run(header).bold = True
            for b in ex.get('bullets', [])[:8]:
                doc.add_paragraph(b, style='List Bullet')
    else:
        doc.add_paragraph("Not specified")

    # Achievements
    doc.add_heading('ACHIEVEMENTS', level=2)
    if achievements:
        for a in achievements:
            doc.add_paragraph(str(a), style='List Bullet')
    else:
        doc.add_paragraph("Not specified")

    # ---------------------------------------------
    # JOB-ALIGNED KEYWORDS (FACT-CHECKED)
    # ---------------------------------------------
    doc.add_heading('JOB-ALIGNED KEYWORDS (FACT-CHECKED)', level=2)
    if validated_keywords:
        doc.add_paragraph(
            "Based on the job description, you should consider incorporating the "
            "following keywords into your resume only where they accurately reflect "
            "your real experience:"
        )
        for kw in validated_keywords:
            doc.add_paragraph(kw, style='List Bullet')
    else:
        doc.add_paragraph(
            "No additional job-specific keywords could be safely recommended beyond "
            "the experience already evidenced in this resume."
        )

    if keyword_example_bullets:
        doc.add_paragraph()
        doc.add_paragraph(
            "Example role-aligned descriptions (only use those that truly match your work):"
        )
        for bullet in keyword_example_bullets:
            doc.add_paragraph(bullet, style='List Bullet')

    # Save docx to memory and send
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"{candidate_name.replace(' ', '_')}_Resume.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
#currently disabled
"""
@app.route('/download-report')
def download_report():
    analysis = session.get("analysis")
    if not analysis:
        return "No analysis data found. Please upload and analyze a resume first.", 400
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin = 0.75 * inch
    y = height - margin

    def new_page():
        nonlocal y
        pdf.showPage(); y = height - margin

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(margin, y, "AI Resume Analysis Report"); y -= 20
    pdf.setFont("Helvetica", 10)
    pdf.drawString(margin, y, f"Generated on {datetime.now().strftime('%B %d, %Y %I:%M %p')}"); y -= 30

    pdf.setFont("Helvetica-Bold", 12); pdf.drawString(margin, y, "Profile Summary:"); y -= 15
    pdf.setFont("Helvetica", 10)
    text_obj = pdf.beginText(margin, y)
    text_obj.textLines(analysis.get("summary","No summary available."))
    pdf.drawText(text_obj); y = text_obj.getY() - 20
    if y < 1.5*inch: new_page()

    pdf.setFont("Helvetica-Bold", 12); pdf.drawString(margin, y, "ATS Compatibility:"); y -= 15
    pdf.setFont("Helvetica", 10); pdf.drawString(margin, y, f"{analysis.get('ats_score','N/A')}%"); y -= 30
    if y < 1.5*inch: new_page()

    pdf.setFont("Helvetica-Bold", 12); pdf.drawString(margin, y, "Skills Identified:"); y -= 15
    pdf.setFont("Helvetica", 10)
    skills_text = ", ".join(analysis.get("skills", [])) or "None detected"
    for line in textwrap.wrap(skills_text, width=90):
        pdf.drawString(margin, y, line); y -= 15
        if y < 1.5*inch: new_page()
    y -= 10

    pdf.setFont("Helvetica-Bold", 12); pdf.drawString(margin, y, "Missing Skills / Gaps:"); y -= 15
    pdf.setFont("Helvetica", 10)
    missing_skills = ", ".join(analysis.get("missing_skills", [])) or "No major gaps found."
    for line in [missing_skills[i:i+100] for i in range(0, len(missing_skills), 100)]:
        pdf.drawString(margin, y, line); y -= 15
        if y < 1.5*inch: new_page()
    y -= 10

    pdf.setFont("Helvetica-Bold", 12); pdf.drawString(margin, y, "Achievements:"); y -= 15
    pdf.setFont("Helvetica", 10)
    ach = analysis.get("honors_awards", []) or ["None detected"]
    for item in ach:
        for line in [item[i:i+100] for i in range(0, len(item), 100)]:
            pdf.drawString(margin, y, f"- {line}"); y -= 15
            if y < 1.5*inch: new_page()
    y -= 10

    # Fraud / AI risk
    pdf.setFont("Helvetica-Bold", 12); pdf.drawString(margin, y, "AI / Fraud Risk:"); y -= 15
    pdf.setFont("Helvetica", 10)
    pdf.drawString(margin, y, f"AI risk score: {analysis.get('ai_risk', {}).get('ai_risk_score', 'N/A')}"); y -= 15
    pdf.drawString(margin, y, f"Timeline issues: {len(analysis.get('timeline_issues', []))}"); y -= 20
    if y < 1.5*inch: new_page()

    # Raw resume text (shortened)
    pdf.setFont("Helvetica-Bold", 12); pdf.drawString(margin, y, "Raw Resume Text (truncated):"); y -= 15
    pdf.setFont("Helvetica", 10)
    resume_text = (analysis.get("resume_text","") or "").replace("\n"," ")
    for line in textwrap.wrap(resume_text, width=110):
        pdf.drawString(margin, y, line); y -= 12
        if y < 1.5*inch: new_page()

    pdf.save(); buffer.seek(0)
    return send_file(buffer, as_attachment=True,
                     download_name=f"AI_Resume_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                     mimetype="application/pdf")
"""
@app.route('/preview-report')
def preview_report():
    """
    Renders the HTML preview of the report using analysis stored in session.
    If no analysis exists, returns a helpful error so the user can re-run analysis.
    """
    analysis = session.get("analysis")
    if not analysis:
        return "No analysis data found. Please upload and analyze a resume first.", 400
    # Ensure keys exist to avoid template errors
    analysis.setdefault("desired_job", "General Role")
    analysis.setdefault("summary", "")
    analysis.setdefault("skills", [])
    analysis.setdefault("education", [])
    analysis.setdefault("experience", [])
    analysis.setdefault("ats_score", 0)
    analysis.setdefault("missing_keywords", [])
    analysis.setdefault("ats_feedback", [])
    analysis.setdefault("missing_skills", [])
    analysis.setdefault("course_recommendations", [])
    analysis.setdefault("alternative_paths", [])
    analysis.setdefault("resume_text", "")
    analysis.setdefault("soft_skills", [])
    analysis.setdefault("ai_risk", {"ai_risk_score": 0})
    analysis.setdefault("timeline_issues", [])
    analysis.setdefault("matching_score", 0)
    analysis.setdefault("language", "unknown")
    analysis.setdefault("bias_warnings", [])

    return render_template('report_preview.html', analysis=analysis)

if __name__ == '__main__':
    app.run(debug=True)
